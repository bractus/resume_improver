import os
from dotenv import load_dotenv
from crewai import Agent, Task, Crew, Process
from crewai.tools import BaseTool
from langchain_openai import ChatOpenAI
from langchain.prompts import PromptTemplate
from crewai_tools import (
    FileReadTool,
    FileWriterTool,
    SerperDevTool
)
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Any, Optional
from pydantic import Field, BaseModel

# Load environment variables
load_dotenv()

# Initialize LLM with fallback options
def get_llm():
    model_name = os.getenv('OPENAI_MODEL_NAME', 'gpt-4')
    try:
        return ChatOpenAI(
            model_name=model_name,
            temperature=0.2
        )
    except:
        # Fallback to GPT-3.5 if GPT-4 is not available
        return ChatOpenAI(
            model_name='gpt-3.5-turbo',
            temperature=0.2
        )

# Initialize LLM
llm = get_llm()

# Set default language
language = "English"

class DocxReaderTool(BaseTool):
    """Tool for reading DOCX files."""
    
    name: str = Field(default="DOCX Reader")
    description: str = Field(default="Read content from a DOCX file")
    
    def __init__(self, file_path: str):
        super().__init__()
        self._file_path = file_path

    def _run(self, *args, **kwargs) -> str:
        """Execute the tool's primary function."""
        try:
            doc = Document(self._file_path)
            full_text = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    full_text.append(paragraph.text.strip())
            return "\n".join(full_text)
        except Exception as e:
            return f"Error reading DOCX file: {str(e)}"

class DocxReaderToolInput(BaseModel):
    file_path: str = Field(description="Path to the DOCX file to read")

class ResumeWriter:
    def __init__(self, output_file="output.docx"):
        self.document = Document()
        self.output_file = output_file
        self.llm = llm
        self._setup_document()
        self._setup_prompts()

    def _setup_document(self):
        # Set up document defaults
        style = self.document.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)

    def _setup_prompts(self):
        self.section_prompt = PromptTemplate(
            input_variables=["section_name", "content"],
            template="""
            Format the following {section_name} section content for a resume, ensuring it is:
            1. Concise and impactful
            2. Uses action verbs
            3. Quantifies achievements where possible
            4. Is ATS-friendly
            5. Maintains professional tone

            Content:
            {content}

            Return only the formatted content without any additional commentary.
            """
        )

    def _format_section_content(self, section_name, content):
        if not content:
            return content

        formatted_content = self.llm.invoke(
            self.section_prompt.format(
                section_name=section_name,
                content=content
            )
        ).content

        return formatted_content

    def write_section(self, title, content):
        # Format content using LLM
        formatted_content = self._format_section_content(title, content)
        
        # Add section title
        heading = self.document.add_paragraph()
        run = heading.add_run(title)
        run.bold = True
        run.font.size = Pt(14)
        
        # Add formatted content
        if isinstance(formatted_content, str):
            paragraphs = formatted_content.split('\n')
            for para in paragraphs:
                if para.strip():
                    if para.strip().startswith('•'):
                        self.document.add_paragraph(
                            para.strip()[1:].strip(),
                            style='List Bullet'
                        )
                    else:
                        self.document.add_paragraph(para.strip())
        elif isinstance(formatted_content, list):
            for item in formatted_content:
                self.document.add_paragraph(item, style='List Bullet')
        
        self.document.add_paragraph()  # Add spacing between sections

    def save(self):
        self.document.save(self.output_file)

# Initialize tools
file_writer_tool = FileWriterTool()
search_tool = SerperDevTool()
read_resume = FileReadTool(file_path='./resume.txt')
example_resume = DocxReaderTool(file_path='./example.docx')

# Agent 1: Resume Researcher
researcher = Agent(
    role="Resume Organizer",
    goal="Create an AI and ATS optimized resume",
    tools=[search_tool, read_resume, file_writer_tool],
    verbose=True,
    backstory=(
        f"As a Resume Organizer, I must create a resume in {language} from the loaded resume by:"
        "- Organizing the resume into sections: Professional Profile, Recent Experience, Education, Skills"
        "- In the Recent Experience section, summarize the latest work experiences"
        "- In the Education section, summarize educational background"
        "- In the Skills section, list skills used in professional projects (from Recent Experience)"
        "- Include certifications in the Professional Profile section"
        "- Include languages in the Skills section"
        "- Aim to fit everything on one page"
        "- Output the content in a structured format that can be parsed into sections"
    ),
    allow_delegation=False
)

# Agent 2: Resume Formatter
formatter = Agent(
    role="Resume Formatter",
    goal="Format the resume to be AI and ATS friendly",
    tools=[example_resume, file_writer_tool],
    backstory=(
        f"As a Resume Formatter, I must format the resume in {language} from the loaded document "
        "and ensure it follows ATS-friendly formatting guidelines:"
        "- Use clear section headings"
        "- Maintain consistent formatting"
        "- Use standard fonts and sizes"
        "- Avoid complex layouts or tables"
        "- Ensure proper spacing between sections"
        "- Follow the example resume format while maintaining ATS compatibility"
    ),
    allow_delegation=False
)

# Agent 3: Manager
manager = Agent(
    role="Manager",
    goal="Efficiently manage the crew and ensure high-quality task completion",
    backstory="You're an experienced project manager, skilled in overseeing complex projects and guiding teams to success. Your role is to coordinate the efforts of the crew members, ensuring that each task is completed on time and to the highest standard.",
    allow_delegation=True,
)

# Define the main task
resume_strategy_task = Task(
    description=(
        f"Create a well-formatted resume in {language}. Structure the content into clear sections "
        "and ensure it follows ATS-friendly formatting guidelines. The output should be organized "
        "into distinct sections that can be properly formatted in a Word document."
    ),
    expected_output=(
        f"A structured resume content that can be formatted into a Microsoft Word document, "
        "with clear sections for Professional Profile, Experience, Education, and Skills."
    ),
    output_file="output.docx"
)

# Create and execute the crew
job_application_crew = Crew(
    agents=[researcher, formatter],
    manager_agent=manager,
    tasks=[resume_strategy_task],
    process=Process.hierarchical,
    verbose=True
)

def format_resume(content):
    writer = ResumeWriter()
    
    # Parse the content and write sections
    sections = {
        "Professional Profile": "",
        "Recent Experience": "",
        "Education": "",
        "Skills": ""
    }
    
    current_section = None
    current_content = []
    
    # Process the content and organize it into sections
    for line in content.split('\n'):
        line = line.strip()
        if line in sections:
            if current_section and current_content:
                sections[current_section] = '\n'.join(current_content)
            current_section = line
            current_content = []
        elif line and current_section:
            current_content.append(line)
    
    # Add the last section
    if current_section and current_content:
        sections[current_section] = '\n'.join(current_content)
    
    # Write sections to document
    for section, content in sections.items():
        if content:
            writer.write_section(section, content)
    
    writer.save()

if __name__ == "__main__":
    # Get the content from the crew
    result = job_application_crew.kickoff()
    
    # Format the result into a proper DOCX file
    format_resume(result)